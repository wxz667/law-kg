from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from dataclasses import dataclass

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from ..utils.ids import checksum_text, slugify
from ..utils.document_layout import (
    APPENDIX_RE,
    ARTICLE_RE,
    ARTICLE_TITLE_SUFFIX_RE,
    CHAPTER_RE,
    PART_RE,
    SECTION_RE,
    is_structural_body_start,
    normalize_heading_text as normalize_heading,
)
from ..utils.numbers import int_to_cn
from ..contracts import LogicalDocumentRecord, PhysicalSourceRecord

OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
DATE_RE = re.compile(r"(\d{4}年\d{1,2}月\d{1,2}日)")
PAREN_DATE_LINE_RE = re.compile(r"^[（(]\d{4}年\d{1,2}月\d{1,2}日.*[）)]$")
EFFECTIVE_DATE_RE = re.compile(r"自(\d{4}年\d{1,2}月\d{1,2}日)起(?:施行|实施|执行)")
REFERENCE_NO_RE = re.compile(
    r"(?:[^\d\s]{1,20})?[〔〔\[]\d{4}[〕\]][^\s]{0,20}?第?\d+号|(?:[^\d\s]{1,20})?[〈<]\d{4}[〉>][^\s]{0,20}?第?\d+号"
)
AUTHORITY_LINE_RE = re.compile(
    r"^(?:中华人民共和国)?(?:.+)?(?:人民法院|人民检察院|人民政府|人民代表大会常务委员会|人民代表大会|办公厅|委员会|部|厅|局)$"
)
ANNOUNCEMENT_RE = re.compile(r"^公\s*告$")
QUOTED_TITLE_RE = re.compile(r"《(?P<title>[^》]{2,200})》")
LOCAL_PCG_RE = re.compile(
    r"(.+?)(?:省|自治区|维吾尔自治区|壮族自治区|回族自治区|特别行政区|市|州|盟|地区|县|区)人民代表大会常务委员会"
)
LOCAL_GOV_RE = re.compile(
    r"(.+?)(?:省|自治区|维吾尔自治区|壮族自治区|回族自治区|特别行政区|市|州|盟|地区|县|区)人民政府"
)
WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

REGION_NAMES = [
    "北京市",
    "天津市",
    "上海市",
    "重庆市",
    "河北省",
    "山西省",
    "辽宁省",
    "吉林省",
    "黑龙江省",
    "江苏省",
    "浙江省",
    "安徽省",
    "福建省",
    "江西省",
    "山东省",
    "河南省",
    "湖北省",
    "湖南省",
    "广东省",
    "海南省",
    "四川省",
    "贵州省",
    "云南省",
    "陕西省",
    "甘肃省",
    "青海省",
    "台湾省",
    "内蒙古自治区",
    "广西壮族自治区",
    "西藏自治区",
    "宁夏回族自治区",
    "新疆维吾尔自治区",
    "香港特别行政区",
    "澳门特别行政区",
]

SOURCE_TYPE_ALIASES = {
    "constitution": "constitution",
    "law": "law",
    "statutes": "law",
    "amendment": "law",
    "regulation": "regulation",
    "decision": "regulation",
    "interpretation": "interpretation",
    "judicial-interpretations": "interpretation",
    "case": "case",
    "cases": "case",
    "宪法": "constitution",
    "法律": "law",
    "行政法规": "regulation",
    "监察法规": "regulation",
    "司法解释": "interpretation",
    "地方性法规": "regulation",
}


class SourceDocumentReadError(ValueError):
    def __init__(self, message: str, *, error_type: str = "source_read_error") -> None:
        super().__init__(message)
        self.error_type = error_type


@dataclass(frozen=True)
class TitleInfo:
    title_lines: list[str]
    consumed_paragraph_count: int


@dataclass(frozen=True)
class TitleBlock:
    title: str
    start: int
    title_end: int
    body_start: int
    issuer: str
    publish_date: str
    reference_no: str
    role: str


def read_source_document(source_path: Path, sidecar_metadata: dict[str, object] | None = None) -> PhysicalSourceRecord:
    validate_source_document(source_path)
    sidecar_metadata = dict(sidecar_metadata) if sidecar_metadata is not None else read_sidecar_metadata(source_path)
    try:
        document = Document(str(source_path))
    except Exception as exc:
        raise SourceDocumentReadError(
            f"Failed to read DOCX package: {exc}",
            error_type="docx_parse_error",
        ) from exc
    raw_paragraphs = extract_document_lines(document)
    paragraphs = normalize_physical_paragraphs(raw_paragraphs)
    title_info = extract_title_info(paragraphs) if paragraphs else TitleInfo([source_path.stem], 1)
    title_lines = title_info.title_lines
    inferred_title = normalize_title_text("".join(title_lines)) if title_lines else normalize_title_text(source_path.stem)
    leading_evidence_text = collect_leading_evidence_text("", paragraphs[:40])
    revision_text = normalize_text("\n".join(paragraphs[:80])).strip()
    revision_events = parse_revision_events(revision_text, title=inferred_title)
    if not revision_events:
        revision_events = parse_revision_events(revision_text or collect_revision_fallback_text(paragraphs, title_info), title=inferred_title)
    title = str(sidecar_metadata.get("title") or inferred_title or source_path.stem)
    source_type = normalize_source_type(
        str(
            sidecar_metadata.get("document_type")
            or sidecar_metadata.get("type")
            or sidecar_metadata.get("source_type")
            or sidecar_metadata.get("category")
            or source_path.parent.name
        )
    )
    issuer = str(
        sidecar_metadata.get("issuer")
        or sidecar_metadata.get("office")
        or infer_issuer(title, revision_text, leading_evidence_text)
    )
    issuer_type = infer_issuer_type(title, revision_text, leading_evidence_text)
    region = infer_region(title, revision_text, leading_evidence_text, issuer_type=issuer_type, source_type=source_type)
    if source_type == "interpretation" and issuer_type in {"supreme_court", "supreme_procuratorate", "joint_judicial", "npc", "npcsc"}:
        region = ""
    document_subtype = infer_document_subtype(source_type, title, revision_text, issuer_type, region, leading_evidence_text)
    source_checksum = checksum_text(normalize_text("\n".join(paragraphs)))
    signature_date = infer_signature_date(paragraphs)
    publish_date = str(
        sidecar_metadata.get("publish_date")
        or sidecar_metadata.get("publish")
        or signature_date
        or (revision_events[0]["date"] if revision_events else "")
    )
    publish_date = normalize_date_value(publish_date)
    effective_date = str(
        sidecar_metadata.get("effective_date")
        or sidecar_metadata.get("expiry")
        or infer_effective_date(revision_text, signature_date, revision_events)
    )
    effective_date = normalize_date_value(effective_date)
    status = normalize_status(sidecar_metadata.get("status"), revision_text)
    source_url = str(sidecar_metadata.get("source_url") or sidecar_metadata.get("detail_url") or sidecar_metadata.get("url") or "")
    base_metadata = {
        "revision_events": revision_events,
        "document_type": source_type,
        "document_subtype": str(sidecar_metadata.get("document_subtype") or document_subtype),
        "issuer": issuer,
        "issuer_type": issuer_type,
        "publish_date": publish_date,
        "effective_date": effective_date,
        "status": status,
        "region": region,
        "source_url": source_url,
        "download_link_word": str(sidecar_metadata.get("download_link_word", "")),
        "download_link_html": str(sidecar_metadata.get("download_link_html", "")),
        "download_link_pdf": str(sidecar_metadata.get("download_link_pdf", "")),
        "crawler_job_id": str(sidecar_metadata.get("crawler_job_id", "")),
    }
    merged_metadata = {key: value for key, value in sidecar_metadata.items() if key != "doc_no"}
    merged_metadata.update(base_metadata)
    return PhysicalSourceRecord(
        source_id=str(sidecar_metadata.get("source_id") or f"{source_type}:{slugify(source_path.stem)}"),
        title=title,
        source_path=str(source_path.resolve()),
        source_type=source_type,
        checksum=str(sidecar_metadata.get("checksum") or source_checksum),
        paragraphs=paragraphs,
        preface_text="",
        toc_lines=[],
        body_lines=paragraphs,
        appendix_lines=[],
        metadata=merged_metadata,
    )


def extract_document_lines(document: DocxDocument) -> list[str]:
    lines: list[str] = []
    numbering_resolver = NumberingResolver(document)
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            text = numbering_resolver.render_paragraph_text(block).strip()
            if text:
                lines.append(text)
            continue
        if isinstance(block, Table):
            lines.extend(render_table_lines(block))
    return lines


def iter_block_items(parent: DocxDocument | _Cell):
    if isinstance(parent, DocxDocument):
        container = parent.element.body
    else:
        container = parent._tc
    for child in container.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def render_table_lines(table: Table) -> list[str]:
    rows = [extract_row_cells(row.cells) for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return []
    if len(rows) == 1:
        return [f"1. {join_labeled_values(rows[0], [])}；"] if any(rows[0]) else []

    header = rows[0]
    body_rows = rows[1:]
    if is_header_row(header, body_rows):
        return [
            f"{index}. {join_labeled_values(row, header)}；"
            for index, row in enumerate(body_rows, start=1)
            if any(row)
        ]
    return [
        f"{index}. {join_labeled_values(row, [])}；"
        for index, row in enumerate(rows, start=1)
        if any(row)
    ]


def extract_row_cells(cells: list[_Cell]) -> list[str]:
    values: list[str] = []
    for cell in cells:
        fragments = [normalize_text(paragraph.text).strip() for paragraph in cell.paragraphs if normalize_text(paragraph.text).strip()]
        values.append(" ".join(fragments).strip())
    return values


def is_header_row(header: list[str], body_rows: list[list[str]]) -> bool:
    if not header or not any(header) or not body_rows:
        return False
    non_empty_headers = [value for value in header if value]
    if len(non_empty_headers) != len(header):
        return False
    if any(len(value) > 20 for value in header):
        return False
    return True


def join_labeled_values(values: list[str], header: list[str]) -> str:
    parts: list[str] = []
    for index, value in enumerate(values):
        clean_value = normalize_text(value).strip()
        if not clean_value:
            continue
        label = header[index].strip() if index < len(header) else ""
        if label:
            parts.append(f"{label}：{clean_value}")
        else:
            parts.append(clean_value)
    return "，".join(parts)


class NumberingResolver:
    def __init__(self, document: DocxDocument) -> None:
        self.levels_by_num_id = load_numbering_levels(document)
        self.counters: dict[int, dict[int, int]] = {}

    def render_paragraph_text(self, paragraph: Paragraph) -> str:
        text = paragraph.text.strip()
        numbering = extract_paragraph_numbering(paragraph)
        if numbering is None:
            return text
        num_id, level = numbering
        level_definition = self.levels_by_num_id.get(num_id, {}).get(level)
        if level_definition is None:
            return text
        counter_state = self.counters.setdefault(num_id, {})
        for deeper_level in list(counter_state):
            if deeper_level > level:
                counter_state.pop(deeper_level, None)
        current_value = counter_state.get(level, int(level_definition.get("start", 1)) - 1) + 1
        counter_state[level] = current_value
        label = render_numbering_label(level_definition["text"], level_definition["formats"], counter_state)
        if not label or text.startswith(label):
            return text
        if text.startswith(("（", "(")) and label.endswith(("）", ")")):
            return text
        return f"{label}{text}"


def extract_paragraph_numbering(paragraph: Paragraph) -> tuple[int, int] | None:
    p_pr = paragraph._p.pPr
    num_pr = getattr(p_pr, "numPr", None) if p_pr is not None else None
    if num_pr is None or num_pr.numId is None or num_pr.ilvl is None:
        return None
    return int(num_pr.numId.val), int(num_pr.ilvl.val)


def load_numbering_levels(document: DocxDocument) -> dict[int, dict[int, dict[str, object]]]:
    try:
        numbering_part = document.part.numbering_part
    except (KeyError, NotImplementedError):
        return {}
    if numbering_part is None:
        return {}
    root = numbering_part._element
    abstract_levels: dict[int, dict[int, dict[str, object]]] = {}
    for abstract_num in root.findall(f"{WORD_NS}abstractNum"):
        abstract_id = int(abstract_num.get(f"{WORD_NS}abstractNumId"))
        abstract_levels[abstract_id] = {}
        for lvl in abstract_num.findall(f"{WORD_NS}lvl"):
            ilvl = int(lvl.get(f"{WORD_NS}ilvl"))
            num_fmt = lvl.find(f"{WORD_NS}numFmt")
            lvl_text = lvl.find(f"{WORD_NS}lvlText")
            start = lvl.find(f"{WORD_NS}start")
            text_pattern = lvl_text.get(f"{WORD_NS}val") if lvl_text is not None else ""
            abstract_levels[abstract_id][ilvl] = {
                "text": text_pattern,
                "start": int(start.get(f"{WORD_NS}val")) if start is not None else 1,
                "formats": collect_level_formats(abstract_num),
            }
            if num_fmt is not None:
                abstract_levels[abstract_id][ilvl]["formats"][ilvl] = num_fmt.get(f"{WORD_NS}val")

    levels_by_num_id: dict[int, dict[int, dict[str, object]]] = {}
    for num in root.findall(f"{WORD_NS}num"):
        num_id = int(num.get(f"{WORD_NS}numId"))
        abstract_ref = num.find(f"{WORD_NS}abstractNumId")
        if abstract_ref is None:
            continue
        abstract_id = int(abstract_ref.get(f"{WORD_NS}val"))
        levels_by_num_id[num_id] = {
            level: dict(definition)
            for level, definition in abstract_levels.get(abstract_id, {}).items()
        }
    return levels_by_num_id


def collect_level_formats(abstract_num) -> dict[int, str]:
    formats: dict[int, str] = {}
    for lvl in abstract_num.findall(f"{WORD_NS}lvl"):
        ilvl = int(lvl.get(f"{WORD_NS}ilvl"))
        num_fmt = lvl.find(f"{WORD_NS}numFmt")
        if num_fmt is not None:
            formats[ilvl] = num_fmt.get(f"{WORD_NS}val")
    return formats


def render_numbering_label(pattern: str, formats: dict[int, str], counters: dict[int, int]) -> str:
    label = pattern
    for level_index, value in counters.items():
        label = label.replace(f"%{level_index + 1}", format_numbering_value(value, formats.get(level_index, "decimal")))
    return label


def format_numbering_value(value: int, fmt: str) -> str:
    if fmt in {"chineseCounting", "chineseCountingThousand"}:
        return int_to_cn(value)
    if fmt in {"decimal", "decimalZero"}:
        return str(value)
    if fmt == "upperLetter":
        return chr(ord("A") + value - 1)
    if fmt == "lowerLetter":
        return chr(ord("a") + value - 1)
    return str(value)


def split_logical_documents(source_document: PhysicalSourceRecord) -> list[LogicalDocumentRecord]:
    paragraphs = [line.strip() for line in (source_document.paragraphs or source_document.body_lines) if line.strip()]
    title_blocks = find_title_blocks(paragraphs, issuer_hint=str(source_document.metadata.get("issuer", "")))
    logical_documents: list[LogicalDocumentRecord] = []
    for index, block in enumerate(title_blocks):
        next_start = title_blocks[index + 1].start if index + 1 < len(title_blocks) else len(paragraphs)
        content_lines = trim_logical_content_lines(
            strip_leading_toc_block(paragraphs[block.body_start:next_start], title=block.title)
        )
        metadata = build_logical_metadata(source_document, block)
        role = str(metadata.get("document_role", ""))
        if role in {"shell", "request"}:
            continue
        body_lines, appendix_lines = split_logical_body_and_appendix(content_lines)
        if not body_lines and not appendix_lines:
            continue
        logical_documents.append(
            LogicalDocumentRecord(
                source_id=build_logical_source_id(source_document, block.title, len(logical_documents) + 1),
                title=block.title,
                source_type=source_document.source_type,
                paragraphs=body_lines,
                appendix_lines=appendix_lines,
                metadata=metadata,
            )
        )
    if logical_documents:
        return logical_documents
    return [build_fallback_logical_document(source_document, paragraphs)]


def normalize_physical_paragraphs(paragraphs: list[str]) -> list[str]:
    normalized: list[str] = []
    for paragraph in paragraphs:
        text = re.sub(r"\s+", " ", paragraph or "").replace("\u3000", " ").strip()
        if not text:
            continue
        if re.match(r"^[\-－—_]+\d+[\-－—_]+$", text):
            continue
        normalized.append(text)
    return normalized


def find_title_blocks(paragraphs: list[str], *, issuer_hint: str = "") -> list[TitleBlock]:
    blocks: list[TitleBlock] = []
    index = 0
    while index < len(paragraphs):
        if not is_title_block_boundary(paragraphs, index):
            index += 1
            continue
        block = extract_title_block_at(paragraphs, index, issuer_hint=issuer_hint)
        if block is None:
            index += 1
            continue
        blocks.append(block)
        index = block.body_start
    return blocks


def is_title_block_boundary(paragraphs: list[str], index: int) -> bool:
    if index <= 0:
        return True
    current = paragraphs[index].strip()
    if not current:
        return False
    previous = ""
    probe = index - 1
    while probe >= 0:
        previous = paragraphs[probe].strip()
        if previous:
            break
        probe -= 1
    if not previous:
        return True
    if is_attachment_marker_line(previous):
        return True
    if DATE_RE.match(previous) or PAREN_DATE_LINE_RE.match(previous):
        return True
    if looks_like_reference_no_line(previous):
        return True
    if normalize_heading(previous) in {"公告", "公 告"}:
        return True
    if any(marker in previous for marker in ("现将", "印发给你们", "转发给你们", "发给你们", "现予公布", "公告公布", "批复如下", "答复如下", "函复如下")):
        return True
    if previous.endswith(("。", "；")) and is_authority_heading_candidate(current):
        return True
    return False


def extract_title_block_at(paragraphs: list[str], start: int, *, issuer_hint: str = "") -> TitleBlock | None:
    if start >= len(paragraphs):
        return None
    first_line = paragraphs[start].strip()
    if not first_line or looks_like_address_line(first_line) or is_structural_body_start(first_line):
        return None

    index = start
    authority_lines: list[str] = []
    while index < len(paragraphs):
        line = paragraphs[index].strip()
        if not line or looks_like_address_line(line):
            return None
        if is_authority_heading_candidate(line, issuer_hint=issuer_hint):
            authority_lines.append(line)
            index += 1
            continue
        break

    title_lines: list[str] = []
    while index < len(paragraphs):
        line = paragraphs[index].strip()
        if not line or looks_like_address_line(line):
            break
        if is_structural_body_start(line) or looks_like_reference_no_line(line) or is_metadata_line(line):
            break
        if normalize_heading(line) in {"公告", "公 告"}:
            return None
        if not looks_like_title_line(line):
            break
        title_lines.append(line)
        index += 1
    if not title_lines:
        return None

    merged_title = normalize_title_text("".join(title_lines))
    if authority_lines and should_prefix_authority_to_title("".join(authority_lines), merged_title):
        merged_title = normalize_title_text("".join(authority_lines) + merged_title)
    allow_generic = start == 0
    if not looks_like_document_title(merged_title, allow_generic=allow_generic):
        return None

    publish_date = ""
    reference_no = ""
    body_start = index
    while body_start < len(paragraphs):
        line = paragraphs[body_start].strip()
        if not line:
            body_start += 1
            continue
        if normalize_heading(line) in {"公告", "公 告"}:
            body_start += 1
            continue
        if DATE_RE.match(line):
            if not publish_date:
                publish_date = normalize_date_value(line)
            body_start += 1
            continue
        if PAREN_DATE_LINE_RE.match(line):
            if not publish_date:
                date_match = DATE_RE.search(line)
                if date_match:
                    publish_date = normalize_date_value(date_match.group(1))
            body_start += 1
            continue
        if looks_like_reference_no_line(line):
            if not reference_no:
                reference_no = line
            body_start += 1
            continue
        break

    issuer = "".join(authority_lines).strip() or str(issuer_hint or "")
    role = classify_document_role(merged_title)
    return TitleBlock(
        title=merged_title,
        start=start,
        title_end=index,
        body_start=body_start,
        issuer=issuer,
        publish_date=publish_date,
        reference_no=reference_no,
        role=role,
    )


def build_logical_metadata(source_document: PhysicalSourceRecord, block: TitleBlock) -> dict[str, object]:
    metadata = dict(source_document.metadata)
    leading_text = normalize_text("\n".join(source_document.paragraphs[max(0, block.start - 2) : block.body_start + 8])).strip()
    issuer = block.issuer or str(metadata.get("issuer") or infer_issuer(block.title, leading_text, leading_text))
    issuer_type = infer_issuer_type(block.title, leading_text, leading_text)
    region = infer_region(
        block.title,
        leading_text,
        leading_text,
        issuer_type=issuer_type,
        source_type=source_document.source_type,
    )
    document_subtype = infer_document_subtype(source_document.source_type, block.title, leading_text, issuer_type, region, leading_text)
    signature_date = block.publish_date or infer_signature_date(source_document.paragraphs[block.start : block.body_start + 20])
    effective_date = infer_effective_date(leading_text, signature_date, [])
    metadata.update(
        {
            "document_type": source_document.source_type,
            "document_subtype": document_subtype or str(metadata.get("document_subtype", "")),
            "issuer": issuer,
            "issuer_type": issuer_type or str(metadata.get("issuer_type", "")),
            "publish_date": normalize_date_value(signature_date or str(metadata.get("publish_date", ""))),
            "effective_date": normalize_date_value(effective_date or str(metadata.get("effective_date", ""))),
            "region": region or str(metadata.get("region", "")),
            "document_role": block.role,
            "reference_no": block.reference_no or str(metadata.get("reference_no", "")),
        }
    )
    return metadata


def build_logical_source_id(source_document: PhysicalSourceRecord, title: str, order: int) -> str:
    base = f"{source_document.source_type}:{slugify(title)}"
    if base != source_document.source_id:
        return base
    return f"{base}:{order:02d}"


def build_fallback_logical_document(source_document: PhysicalSourceRecord, paragraphs: list[str]) -> LogicalDocumentRecord:
    metadata = dict(source_document.metadata)
    metadata.setdefault("document_type", source_document.source_type)
    metadata.setdefault("document_role", "substantive")
    filtered_paragraphs = strip_leading_toc_block(paragraphs, title=source_document.title)
    body_lines, appendix_lines = split_logical_body_and_appendix(filtered_paragraphs)
    return LogicalDocumentRecord(
        source_id=source_document.source_id,
        title=source_document.title,
        source_type=source_document.source_type,
        paragraphs=body_lines,
        appendix_lines=appendix_lines,
        metadata=metadata,
    )


def trim_logical_content_lines(lines: list[str]) -> list[str]:
    trimmed = [line.strip() for line in lines if line.strip()]
    while trimmed and is_attachment_marker_line(trimmed[-1]):
        trimmed.pop()
    return trimmed


def strip_leading_toc_block(lines: list[str], *, title: str) -> list[str]:
    normalized_lines = [line.strip() for line in lines if line.strip()]
    if not normalized_lines:
        return []
    toc_index = next(
        (index for index, line in enumerate(normalized_lines[:20]) if normalize_heading(line) == "目录"),
        None,
    )
    if toc_index is None:
        return normalized_lines
    synthetic_paragraphs = [title.strip(), *normalized_lines]
    _, _, content_lines, appendix_lines = split_document_sections(
        synthetic_paragraphs,
        TitleInfo([title.strip()], 1),
    )
    filtered_lines = [line.strip() for line in [*content_lines, *appendix_lines] if line.strip()]
    return filtered_lines or normalized_lines


def split_logical_body_and_appendix(lines: list[str]) -> tuple[list[str], list[str]]:
    normalized = [line.strip() for line in lines if line.strip()]
    if not normalized:
        return [], []
    appendix_index = next((idx for idx, line in enumerate(normalized) if APPENDIX_RE.match(normalize_heading(line))), None)
    if appendix_index is None:
        return normalized, []
    return normalized[:appendix_index], normalized[appendix_index:]


def looks_like_document_title(title: str, *, allow_generic: bool) -> bool:
    normalized = normalize_title_text(title)
    if not normalized or len(normalized) < 2:
        return False
    if normalized in {"公告", "公 告"}:
        return False
    if is_structural_body_start(normalized):
        return False
    if any(marker in normalized for marker in ("：", "。", "；")):
        return False
    if any(keyword in normalized for keyword in ("现将", "请遵照", "请参照", "请认真", "通知如下")):
        return False
    if any(normalized.endswith(suffix) for suffix in ("意见", "意见（试行）", "规定", "办法", "解释", "细则", "规则", "通则", "决定", "条例", "批复", "答复", "复函", "通知", "公告", "请示", "报告")):
        return True
    if "关于" in normalized and any(keyword in normalized for keyword in ("意见", "规定", "办法", "解释", "批复", "答复", "通知", "请示", "报告")):
        return True
    return allow_generic and len(normalized) <= 40


def should_prefix_authority_to_title(authority_text: str, title: str) -> bool:
    normalized_authority = normalize_heading(authority_text)
    normalized_title = normalize_title_text(title)
    if not normalized_authority or not normalized_title or normalized_title.startswith(normalized_authority):
        return False
    if normalized_title.startswith("关于《") and normalized_title.endswith("的解释"):
        return True
    return False


def classify_document_role(title: str) -> str:
    normalized = normalize_title_text(title)
    if any(keyword in normalized for keyword in ("批复", "答复", "复函", "函复")):
        return "reply"
    if any(keyword in normalized for keyword in ("请示", "报告")):
        return "request"
    if any(keyword in normalized for keyword in ("通知", "公告", "印发", "转发", "通报")):
        return "shell"
    if any(normalized.endswith(suffix) for suffix in ("意见", "意见（试行）", "规定", "办法", "解释", "细则", "规则", "通则", "决定", "条例")):
        return "substantive"
    return "substantive"


def is_metadata_line(line: str) -> bool:
    return bool(DATE_RE.match(line) or PAREN_DATE_LINE_RE.match(line) or looks_like_reference_no_line(line))


def is_attachment_marker_line(line: str) -> bool:
    return bool(re.match(r"^附(?:件)?(?:[：:]|[一二三四五六七八九十百千万零两〇0-9]+)?$", normalize_heading(line)))


def looks_like_title_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.endswith("："):
        return False
    if looks_like_reference_no_line(stripped):
        return False
    if stripped in {"序号", "名称", "司法解释名称", "发文日期、文号", "废止理由", "理由"}:
        return False
    if re.match(r"^[0-9０-９]+$", stripped):
        return False
    if is_structural_body_start(stripped):
        return False
    if re.match(r"^[一二三四五六七八九十百千]+、.+$", stripped):
        return False
    if re.match(r"^(?:（|\()[一二三四五六七八九十]+(?:）|\)).+$", stripped):
        return False
    if re.match(r"^(?:[0-9０-９]+[.．、]|(?:（|\()[0-9０-９]+(?:）|\))).+$", stripped):
        return False
    if any(token in stripped for token in ("。", "；")):
        return False
    return True


def looks_like_address_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped.endswith("："):
        return False
    normalized = normalize_heading(stripped)
    return any(keyword in normalized for keyword in ("人民法院", "人民检察院", "各省", "各自治区", "各直辖市", "各部门", "各单位", "最高人民法院", "最高人民检察院"))


def is_authority_heading_candidate(line: str, *, issuer_hint: str = "") -> bool:
    stripped = line.strip()
    if not stripped or stripped.endswith("："):
        return False
    normalized = normalize_heading(stripped)
    if issuer_hint and normalized in {normalize_heading(issuer_hint), f"中华人民共和国{normalize_heading(issuer_hint)}"}:
        return True
    return bool(AUTHORITY_LINE_RE.match(normalized))


def read_sidecar_metadata(source_path: Path) -> dict[str, object]:
    sidecar_path = source_path.with_suffix(".metadata.json")
    if not sidecar_path.exists():
        return {}
    try:
        payload = json.loads(sidecar_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise SourceDocumentReadError(
            f"Invalid source metadata JSON: {sidecar_path} ({exc})",
            error_type="invalid_metadata_json",
        ) from exc
    if not isinstance(payload, dict):
        raise SourceDocumentReadError(
            f"Source metadata must be an object: {sidecar_path}",
            error_type="invalid_metadata_json",
        )
    return payload


def validate_source_document(source_path: Path) -> None:
    if not source_path.exists():
        raise SourceDocumentReadError(
            f"Source file not found: {source_path}",
            error_type="missing_source_file",
        )
    if source_path.suffix.lower() != ".docx":
        raise SourceDocumentReadError(
            f"Unsupported source format: {source_path.suffix}",
            error_type="unsupported_source_format",
        )
    try:
        file_header = source_path.read_bytes()[:8]
    except OSError as exc:
        raise SourceDocumentReadError(
            f"Failed to read source file header: {exc}",
            error_type="source_read_error",
        ) from exc
    if file_header == OLE2_MAGIC:
        raise SourceDocumentReadError(
            "Source file is a legacy OLE Word document mislabeled as .docx; "
            "please convert it to a real DOCX package before building.",
            error_type="legacy_word_format",
        )
    if not zipfile.is_zipfile(source_path):
        raise SourceDocumentReadError(
            f"Source file is not a valid DOCX zip package: {source_path}",
            error_type="invalid_docx_package",
        )

    try:
        with zipfile.ZipFile(source_path) as archive:
            names = set(archive.namelist())
    except zipfile.BadZipFile as exc:
        raise SourceDocumentReadError(
            f"Broken DOCX zip package: {exc}",
            error_type="invalid_docx_package",
        ) from exc

    required_entries = {
        "[Content_Types].xml",
        "_rels/.rels",
        "word/document.xml",
    }
    missing_entries = sorted(required_entries - names)
    if missing_entries:
        raise SourceDocumentReadError(
            "DOCX package is missing required entries: " + ", ".join(missing_entries),
            error_type="incomplete_docx_package",
        )


def normalize_source_type(raw_type: str) -> str:
    return SOURCE_TYPE_ALIASES.get(raw_type, slugify(raw_type))


def normalize_status(raw_status: object, preface_text: str) -> str:
    text = str(raw_status or "").strip().lower()
    if text in {"1", "有效", "现行有效", "effective"}:
        return "effective"
    if text in {"已修改", "修正", "修订", "amended"}:
        return "amended"
    if text in {"废止", "失效", "0", "repealed", "invalid"}:
        return "repealed"
    inferred = infer_status(preface_text)
    return inferred if inferred else "unknown"


def infer_issuer_type(title: str, preface_text: str, leading_evidence_text: str = "") -> str:
    text = "\n".join(part for part in (title, preface_text, leading_evidence_text) if part).strip()
    authority_lines = collect_authority_lines(title, preface_text, leading_evidence_text)
    if matches_authority(authority_lines, text, "最高人民法院") and matches_authority(authority_lines, text, "最高人民检察院"):
        return "joint_judicial"
    if matches_authority(authority_lines, text, "最高人民法院"):
        return "supreme_court"
    if matches_authority(authority_lines, text, "最高人民检察院"):
        return "supreme_procuratorate"
    if matches_authority(authority_lines, text, "国家监察委员会"):
        return "supervisory_commission"
    if matches_authority(authority_lines, text, "全国人民代表大会常务委员会"):
        return "npcsc"
    if matches_authority(authority_lines, text, "全国人民代表大会"):
        return "npc"
    if matches_authority(authority_lines, text, "国务院"):
        return "state_council"
    if LOCAL_PCG_RE.search(text):
        return "local_people_congress"
    if LOCAL_GOV_RE.search(text):
        return "local_government"
    if re.search(r"(?:省|自治区|直辖市|自治州|地区|市|县|区).{0,10}人民检察院", text):
        return "local_procuratorate"
    if re.search(r"(?:省|自治区|直辖市|自治州|地区|市|县|区).{0,10}人民法院", text):
        return "local_court"
    return ""


def infer_issuer(title: str, preface_text: str, leading_evidence_text: str = "") -> str:
    text = "\n".join(part for part in (title, preface_text, leading_evidence_text) if part).strip()
    authority_lines = collect_authority_lines(title, preface_text, leading_evidence_text)
    for candidate in (
        "最高人民法院、最高人民检察院",
        "最高人民检察院",
        "最高人民法院",
        "全国人民代表大会常务委员会",
        "全国人民代表大会",
        "国家监察委员会",
        "国务院",
    ):
        if matches_authority(authority_lines, text, candidate):
            return candidate
    local_match = LOCAL_PCG_RE.search(text) or LOCAL_GOV_RE.search(text)
    if local_match:
        return local_match.group(0)
    local_procuratorate_match = re.search(r"(?:省|自治区|直辖市|自治州|地区|市|县|区).{0,10}人民检察院", text)
    if local_procuratorate_match:
        return local_procuratorate_match.group(0)
    local_court_match = re.search(r"(?:省|自治区|直辖市|自治州|地区|市|县|区).{0,10}人民法院", text)
    if local_court_match:
        return local_court_match.group(0)
    return ""


def collect_authority_lines(title: str, preface_text: str, leading_evidence_text: str = "") -> list[str]:
    raw_lines = [
        line.strip()
        for line in "\n".join(part for part in (title, preface_text, leading_evidence_text) if part).splitlines()
        if line.strip()
    ]
    return raw_lines[:12]


def has_authority(lines: list[str], candidate: str) -> bool:
    variants = {candidate, f"中华人民共和国{candidate}"}
    normalized_lines = [normalize_heading(line) for line in lines]
    normalized_variants = {normalize_heading(item) for item in variants}
    return any(line in normalized_variants for line in normalized_lines)


def matches_authority(lines: list[str], text: str, candidate: str) -> bool:
    return has_authority(lines, candidate) or candidate in text or f"中华人民共和国{candidate}" in text


def infer_region(
    title: str,
    preface_text: str,
    leading_evidence_text: str = "",
    *,
    issuer_type: str = "",
    source_type: str = "",
) -> str:
    if issuer_type in {
        "supreme_court",
        "supreme_procuratorate",
        "joint_judicial",
        "supervisory_commission",
        "npc",
        "npcsc",
        "state_council",
    }:
        return ""
    if source_type == "interpretation" and issuer_type in {"local_court", "local_procuratorate"}:
        text = "\n".join(part for part in (title, preface_text, leading_evidence_text) if part).strip()
        for region_name in REGION_NAMES:
            if region_name in text:
                return region_name
        return ""
    text = "\n".join(part for part in (title, preface_text, leading_evidence_text) if part).strip()
    for region_name in REGION_NAMES:
        if region_name in text:
            return region_name
    return ""


def infer_document_subtype(
    source_type: str,
    title: str,
    preface_text: str,
    issuer_type: str,
    region: str,
    leading_evidence_text: str = "",
) -> str:
    evidence_text = "\n".join(part for part in (title, preface_text, leading_evidence_text) if part)
    normalized_title = normalize_title_text(title)
    if source_type == "constitution":
        if "修正案" in title or "修正文本" in title:
            return "amendment"
        return ""
    if source_type == "law":
        if "修正案" in evidence_text or "修订" in evidence_text or "修正" in evidence_text:
            return "amendment"
        return ""
    if source_type == "regulation":
        if title_indicates_decision(normalized_title):
            return "decision"
        if issuer_type == "state_council":
            return "administrative"
        if issuer_type == "supervisory_commission":
            return "supervisory"
        if region or issuer_type in {"local_people_congress", "local_government"}:
            return "local"
        return ""
    if source_type == "interpretation":
        if issuer_type in {"npc", "npcsc"}:
            return "legislative"
        if issuer_type in {"supreme_court", "supreme_procuratorate", "joint_judicial"}:
            return "judicial"
        if "决定" in evidence_text:
            return "decision"
        return ""
    return ""


def collect_leading_evidence_text(preface_text: str, body_lines: list[str]) -> str:
    body_head = "\n".join(line.strip() for line in body_lines[:16] if line.strip())
    return "\n".join(part for part in (preface_text, body_head) if part).strip()


def collect_revision_fallback_text(paragraphs: list[str], title_info: TitleInfo) -> str:
    lines: list[str] = []
    start = title_info.consumed_paragraph_count
    for line in paragraphs[start : start + 12]:
        stripped = line.strip()
        if not stripped:
            continue
        normalized = normalize_heading(stripped)
        if normalized.startswith("附件"):
            break
        if stripped.endswith("："):
            break
        if is_structural_body_start(stripped):
            break
        if re.match(r"^(?:[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）)", stripped):
            break
        lines.append(stripped)
    return normalize_text("\n".join(lines)).strip()


def infer_status(preface_text: str) -> str:
    text = preface_text or ""
    if any(marker in text for marker in ("予以废止", "同时废止", "已经废止", "宣布废止", "废止。", "废止；", "废止,", "废止，")):
        return "repealed"
    if "修正" in text or "修订" in text:
        return "amended"
    return "effective"


def title_indicates_decision(title: str) -> bool:
    normalized = normalize_title_text(title)
    if not normalized.endswith("决定"):
        return False
    return not normalized.endswith(("条例", "规定", "办法", "规则", "细则", "通则"))


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
    normalized = "\n".join(line for line in lines if line)
    return normalized.strip() + ("\n" if normalized else "")


def normalize_title_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip()


def extract_title(paragraphs: list[str]) -> str:
    return normalize_title_text("".join(extract_title_info(paragraphs).title_lines))


def extract_title_lines(paragraphs: list[str]) -> list[str]:
    return extract_title_info(paragraphs).title_lines


def extract_title_info(paragraphs: list[str]) -> TitleInfo:
    if not paragraphs:
        return TitleInfo([], 0)
    announcement_title = extract_announcement_title(paragraphs)
    if announcement_title:
        return TitleInfo([announcement_title], 0)
    title_lines: list[str] = []
    for line in paragraphs:
        stripped = line.strip()
        if not stripped:
            if title_lines:
                break
            continue
        if title_lines and _is_title_terminator(stripped):
            break
        if not title_lines and normalize_heading(stripped) == "目录":
            break
        title_lines.append(stripped)
    if not title_lines:
        return TitleInfo([paragraphs[0].strip()], 1)
    return TitleInfo(title_lines, len(title_lines))


def extract_announcement_title(paragraphs: list[str]) -> str:
    normalized = [normalize_heading(item) for item in paragraphs[:8]]
    has_announcement = any(ANNOUNCEMENT_RE.match(item) for item in normalized)
    if not has_announcement:
        return ""
    for line in paragraphs[:8]:
        match = QUOTED_TITLE_RE.search(line)
        if not match:
            continue
        if any(keyword in line for keyword in ("现予公布", "已于", "施行")):
            return match.group("title").strip()
    return ""


def _is_title_terminator(line: str) -> bool:
    normalized = normalize_heading(line)
    if normalized == "目录":
        return True
    if normalized in {"序号", "名称", "司法解释名称", "发文日期、文号", "废止理由"}:
        return True
    if normalized in {"序言", "前言"}:
        return True
    if QUOTED_TITLE_RE.search(line) and any(keyword in line for keyword in ("已于", "现予", "公布", "施行", "发布")):
        return True
    if line.endswith("：") and any(keyword in normalized for keyword in ("人民检察院", "人民法院", "各省", "各自治区", "各直辖市", "各部门", "各单位")):
        return True
    if ARTICLE_TITLE_SUFFIX_RE.match(line):
        return False
    if is_structural_body_start(line):
        return True
    if re.match(r"^(?:[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）)", line.strip()):
        return True
    if any(keyword in normalized for keyword in ("根据", "为了", "现就", "解释如下", "作如下规定", "通知如下")):
        return True
    if re.match(r"^[（(]\d{4}年\d{1,2}月\d{1,2}日", line):
        return True
    if re.match(r"^\d{4}年\d{1,2}月\d{1,2}日", line):
        return True
    if looks_like_reference_no_line(line):
        return True
    return False


def infer_signature_date(paragraphs: list[str]) -> str:
    standalone_dates = [
        line.strip()
        for line in paragraphs[:40]
        if re.match(r"^\d{4}年\d{1,2}月\d{1,2}日$", line.strip())
    ]
    return standalone_dates[-1] if standalone_dates else ""


def infer_effective_date(preface_text: str, signature_date: str, revision_events: list[dict[str, str]]) -> str:
    explicit_match = EFFECTIVE_DATE_RE.search(preface_text)
    if explicit_match:
        return explicit_match.group(1)
    if "自即日起" in preface_text and signature_date:
        return signature_date
    if revision_events:
        return revision_events[-1]["date"]
    return signature_date


def normalize_date_value(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    iso_match = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})$", text)
    if iso_match:
        year, month, day = iso_match.groups()
        return f"{year}-{int(month):02d}-{int(day):02d}"
    cn_match = re.match(r"^(\d{4})年(\d{1,2})月(\d{1,2})日$", text)
    if cn_match:
        year, month, day = cn_match.groups()
        return f"{year}-{int(month):02d}-{int(day):02d}"
    return text


def split_document_sections(
    paragraphs: list[str],
    title_info: TitleInfo,
) -> tuple[str, list[str], list[str], list[str]]:
    if not paragraphs:
        return "", [], [], []
    title_lines = title_info.title_lines
    title_line_count = title_info.consumed_paragraph_count

    toc_index = next((index for index, line in enumerate(paragraphs) if normalize_heading(line) == "目录"), None)
    toc_lines: list[str] = []
    content_lines: list[str] = paragraphs[1:]
    body_start = 1
    if toc_index is not None:
        body_start = find_body_start_index(paragraphs, toc_index, title_lines, title_line_count)
        toc_lines = paragraphs[toc_index + 1 : body_start]
        content_lines = paragraphs[body_start:]
        preface_lines = paragraphs[title_line_count:toc_index]
    else:
        body_start = find_body_start_without_toc(paragraphs, title_lines, title_line_count)
        preface_lines = paragraphs[title_line_count:body_start]
        content_lines = paragraphs[body_start:]

    preface_text = normalize_text("\n".join(preface_lines)).strip()

    appendix_start = next(
        (index for index, line in enumerate(content_lines) if APPENDIX_RE.match(line)),
        None,
    )
    if appendix_start is None:
        return preface_text, toc_lines, content_lines, []
    return (
        preface_text,
        toc_lines,
        content_lines[:appendix_start],
        content_lines[appendix_start:],
    )


def find_body_start_index(paragraphs: list[str], toc_index: int, title_lines: list[str], title_line_count: int) -> int:
    repeated_title_index = find_repeated_title_index(paragraphs, title_lines, title_line_count)
    if repeated_title_index is not None and repeated_title_index > toc_index:
        search_start = skip_title_block(paragraphs, repeated_title_index, title_lines)
        return skip_cover_block(paragraphs, search_start)
    part_indices = [index for index, line in enumerate(paragraphs) if PART_RE.match(line)]
    if len(part_indices) >= 2:
        first_part_line = paragraphs[part_indices[0]]
        for later_index in part_indices[1:]:
            if paragraphs[later_index] == first_part_line and later_index > toc_index:
                return later_index
    repeated_toc_heading_index = find_repeated_toc_heading_index(paragraphs, toc_index)
    if repeated_toc_heading_index is not None:
        return repeated_toc_heading_index

    first_structural = next(
        (
            index
            for index, line in enumerate(paragraphs[toc_index + 1 :], start=toc_index + 1)
            if is_structural_body_start(line)
        ),
        None,
    )
    if first_structural is not None:
        return first_structural
    return toc_index + 1


def find_body_start_without_toc(paragraphs: list[str], title_lines: list[str], title_line_count: int) -> int:
    repeated_title_index = find_repeated_title_index(paragraphs, title_lines, title_line_count)
    if repeated_title_index is not None:
        search_start = skip_title_block(paragraphs, repeated_title_index, title_lines)
        return skip_cover_block(paragraphs, search_start)
    first_structural = next(
        (
            index
            for index, line in enumerate(paragraphs[title_line_count:], start=title_line_count)
            if is_structural_body_start(line)
        ),
        None,
    )
    if first_structural is not None:
        return first_structural

    index = title_line_count
    while index < len(paragraphs):
        line = paragraphs[index].strip()
        if not line:
            index += 1
            continue
        if re.match(r"^[（(]\d{4}年\d{1,2}月\d{1,2}日.*[）)]$", line):
            index += 1
            continue
        if re.match(r"^\d{4}年\d{1,2}月\d{1,2}日", line):
            index += 1
            continue
        break
    return index if index < len(paragraphs) else title_line_count


def find_repeated_title_index(paragraphs: list[str], title_lines: list[str], title_line_count: int) -> int | None:
    if not paragraphs or not title_lines:
        return None
    title = normalize_heading("".join(title_lines))
    for index, line in enumerate(paragraphs[title_line_count:], start=title_line_count):
        if normalize_heading(line) == title:
            return index
    span = len(title_lines)
    if span <= 1:
        return None
    for index in range(title_line_count, len(paragraphs) - span + 1):
        merged = "".join(normalize_heading(paragraphs[index + offset]) for offset in range(span))
        if merged == title:
            return index
    return None


def skip_title_block(paragraphs: list[str], start_index: int, title_lines: list[str]) -> int:
    joined_title = normalize_heading("".join(title_lines))
    if start_index >= len(paragraphs):
        return start_index
    if normalize_heading(paragraphs[start_index]) == joined_title:
        return start_index + 1
    span = len(title_lines)
    if span and start_index + span <= len(paragraphs):
        merged = "".join(normalize_heading(paragraphs[start_index + offset]) for offset in range(span))
        if merged == joined_title:
            return start_index + span
    return start_index


def skip_cover_block(paragraphs: list[str], start_index: int) -> int:
    index = start_index
    while index < len(paragraphs):
        line = paragraphs[index].strip()
        if not line:
            index += 1
            continue
        normalized = normalize_heading(line)
        if looks_like_reference_no_line(normalized):
            index += 1
            continue
        if ANNOUNCEMENT_RE.match(normalized):
            index += 1
            continue
        if line.startswith(("（", "(")) and re.search(r"\d{4}年\d{1,2}月\d{1,2}日", line):
            index += 1
            while index < len(paragraphs):
                follow = paragraphs[index].strip()
                if not follow:
                    index += 1
                    continue
                if looks_like_reference_no_line(follow):
                    index += 1
                    continue
                if follow.endswith(("）", ")")):
                    index += 1
                break
            continue
        if re.match(r"^\d{4}年\d{1,2}月\d{1,2}日$", line):
            index += 1
            continue
        break
    return index if index < len(paragraphs) else start_index


def find_repeated_toc_heading_index(paragraphs: list[str], toc_index: int) -> int | None:
    toc_structural_lines = [
        normalize_heading(line)
        for line in paragraphs[toc_index + 1 :]
        if is_structural_body_start(line)
    ]
    seen: set[str] = set()
    unique_toc_lines: list[str] = []
    for line in toc_structural_lines:
        if line in seen:
            continue
        seen.add(line)
        unique_toc_lines.append(line)
        if len(unique_toc_lines) >= 8:
            break
    if not unique_toc_lines:
        return None
    for index in range(toc_index + 1, len(paragraphs)):
        normalized = normalize_heading(paragraphs[index])
        if normalized in unique_toc_lines and any(
            normalize_heading(paragraphs[later]) == normalized
            for later in range(index + 1, len(paragraphs))
        ):
            continue
        if normalized in unique_toc_lines:
            return index
    return None


def parse_revision_events(preface_text: str, *, title: str = "") -> list[dict[str, str]]:
    events: list[dict[str, str]] = []
    if not preface_text:
        return events
    chunks = build_revision_event_chunks(preface_text)
    for chunk in chunks:
        for date_text, description in extract_revision_event_sentences(chunk, title=title):
            events.append(
                {
                    "date": normalize_date_value(date_text),
                    "event_type": classify_revision_event(description),
                    "description": description,
                }
            )
    return events


def build_revision_event_chunks(preface_text: str) -> list[str]:
    cleaned_lines = []
    for raw_line in preface_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if normalize_heading(line).startswith("附件"):
            break
        if looks_like_reference_no_line(line):
            continue
        cleaned_lines.append(line)
    chunks: list[str] = []
    current: list[str] = []
    for line in cleaned_lines:
        current.append(line)
        if any(symbol in line for symbol in "。；;）)"):
            chunk = " ".join(current).strip()
            if DATE_RE.search(chunk):
                chunks.append(chunk)
            current = []
    if current:
        chunk = " ".join(current).strip()
        if DATE_RE.search(chunk):
            chunks.append(chunk)
    return chunks


def extract_revision_event_sentences(chunk: str, *, title: str = "") -> list[tuple[str, str]]:
    cleaned_chunk = clean_revision_event_description(chunk)
    if not cleaned_chunk:
        return []
    if should_skip_revision_event_chunk(cleaned_chunk, title=title):
        return []
    amendment_clauses = split_based_on_amendment_clauses(cleaned_chunk)
    if amendment_clauses:
        return amendment_clauses
    sequential_clauses = split_sequential_event_clauses(cleaned_chunk, title=title)
    if sequential_clauses:
        return sequential_clauses
    matches = list(DATE_RE.finditer(cleaned_chunk))
    if not matches:
        return []
    if len(matches) == 1:
        return [(matches[0].group(1), cleaned_chunk)]
    if all(keyword not in cleaned_chunk for keyword in ("公布", "施行")):
        clauses: list[tuple[str, str]] = []
        for index, match in enumerate(matches):
            start = match.start()
            end = matches[index + 1].start() if index + 1 < len(matches) else len(cleaned_chunk)
            clause = clean_revision_event_description(cleaned_chunk[start:end].strip(" ，；;"))
            if not clause or should_skip_revision_event_chunk(clause, title=title):
                continue
            if not any(keyword in clause for keyword in ("通过", "修订", "修正", "施行", "公布")):
                continue
            clauses.append((match.group(1), clause))
        if clauses:
            return clauses
    return [(match.group(1), cleaned_chunk) for match in matches]


def split_sequential_event_clauses(text: str, *, title: str = "") -> list[tuple[str, str]]:
    matches = list(DATE_RE.finditer(text))
    if len(matches) < 2:
        return []
    clauses: list[tuple[str, str]] = []
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        clause = clean_revision_event_description(text[start:end].strip(" ，；;"))
        clause = re.sub(r"[，,]?\s*自$", "", clause).strip()
        if not clause or should_skip_revision_event_chunk(clause, title=title):
            continue
        if not any(keyword in clause for keyword in ("通过", "修订", "修正", "施行", "公布")):
            continue
        clauses.append((match.group(1), clause))
    return clauses


def split_based_on_amendment_clauses(text: str) -> list[tuple[str, str]]:
    if "根据" not in text or "修正" not in text:
        return []
    _, _, tail = text.partition("根据")
    if not tail:
        return []
    tail = tail.strip().rstrip("）)")
    if not tail.endswith("修正"):
        return []
    amendment_list = tail[: -len("修正")].strip(" ，、")
    if not amendment_list:
        return []
    parts = [
        clean_revision_event_description(part.strip(" ，；;"))
        for part in re.split(r"、|和(?=\d{4}年)", amendment_list)
        if part.strip(" ，；;")
    ]
    clauses: list[tuple[str, str]] = []
    for part in parts:
        match = DATE_RE.search(part)
        if not match:
            continue
        description = f"{part}修正"
        if "修正案" not in description and "修正" not in description:
            continue
        clauses.append((match.group(1), description))
    return clauses


def clean_revision_event_description(text: str) -> str:
    text = text.strip()
    if (text.startswith("（") and text.endswith("）")) or (text.startswith("(") and text.endswith(")")):
        text = text[1:-1].strip()
    text = text.rstrip("）)")
    text = re.sub(r"\s+", " ", text)
    return text


def looks_like_reference_no_line(text: str) -> bool:
    if not text:
        return False
    normalized = normalize_heading(text.strip())
    normalized = re.sub(r"^\d{4}年\d{1,2}月\d{1,2}日", "", normalized)
    if normalized.startswith(("你院", "你部", "你厅", "你局", "你省", "你市", "你县", "你庭", "你所")):
        return False
    return bool(REFERENCE_NO_RE.search(normalized))


def should_skip_revision_event_chunk(text: str, *, title: str = "") -> bool:
    if not text:
        return True
    normalized_title = normalize_heading(title)
    normalized_text = normalize_heading(text)
    if (
        normalized_title
        and normalized_title in normalized_text
        and "通过" in text
        and not any(keyword in text for keyword in ("公布", "施行", "修订", "修正"))
    ):
        return True
    if any(keyword in text for keyword in ("通过", "修订", "修正", "施行", "公布")):
        return False
    compact = normalize_heading(text)
    compact = DATE_RE.sub("", compact)
    compact = REFERENCE_NO_RE.sub("", compact)
    for authority in (
        "中华人民共和国最高人民检察院",
        "中华人民共和国最高人民法院",
        "中华人民共和国全国人民代表大会常务委员会",
        "最高人民检察院",
        "最高人民法院",
        "全国人民代表大会常务委员会",
        "公告",
    ):
        compact = compact.replace(normalize_heading(authority), "")
    compact = compact.replace(normalized_title, "")
    return not compact


def classify_revision_event(description: str) -> str:
    if "修订" in description:
        return "revision"
    if "修正" in description or "修正案" in description:
        return "amendment"
    if any(keyword in description for keyword in ("通过", "公布", "施行")):
        return "adoption"
    return "other"
