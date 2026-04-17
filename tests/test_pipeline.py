from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from builder.cli import StageBarDisplay, build_parser, split_graph_export
from builder.io import read_graph_bundle, read_manifest, read_normalize_index
from builder.pipeline.orchestrator import build_batch_knowledge_graph, build_knowledge_graph, resolve_source_id
from builder.io import read_reference_candidates, read_relation_plans
from builder.stages.reference_filter.helpers import build_relation_context


def write_pipeline_source(
    root: Path,
    *,
    source_id: str = "law:sample",
    title: str = "示例法",
    category: str = "法律",
    issuer: str = "全国人民代表大会常务委员会",
    status: str = "现行有效",
) -> str:
    source_dir = root / "source" / "docs"
    source_dir.mkdir(parents=True, exist_ok=True)
    source_path = source_dir / f"{title}.docx"
    document = Document()
    document.add_paragraph(title)
    document.add_paragraph("第一条 机动车，是指以动力装置驱动的车辆。")
    document.add_paragraph("第二条 依照《示例法》第一条，机动车发生交通事故的，依照本条处理。")
    document.add_paragraph("第三条 前款规定的交通事故，应当依法处理。")
    document.save(str(source_path))

    metadata_dir = root / "source" / "metadata"
    metadata_dir.mkdir(parents=True, exist_ok=True)
    metadata_path = metadata_dir / "metadata-0001.json"
    payload = []
    if metadata_path.exists():
        payload = json.loads(metadata_path.read_text(encoding="utf-8"))
    payload.append(
        {
            "source_id": source_id,
            "title": title,
            "issuer": issuer,
            "publish_date": "2024-01-01",
            "document_type": "law",
            "category": category,
            "status": status,
            "source_format": "docx",
        }
    )
    metadata_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return source_id


def test_pipeline_builds_final_graph_bundle(tmp_path: Path) -> None:
    source_id = write_pipeline_source(tmp_path)
    result = build_knowledge_graph(
        source_id=source_id,
        data_root=tmp_path,
        through_stage="implicit_reasoning",
    )
    final_graph_path = Path(result["artifact_paths"]["final_graph_bundle"])
    bundle = read_graph_bundle(final_graph_path)

    assert bundle.metadata["stage"] == "implicit_reasoning"
    assert any(node.level == "concept" for node in bundle.nodes)
    assert any(edge.type in {"REFERENCES", "INTERPRETS"} for edge in bundle.edges)


def test_split_export_writes_jsonl(tmp_path: Path) -> None:
    source_id = write_pipeline_source(tmp_path)
    result = build_knowledge_graph(
        source_id=source_id,
        data_root=tmp_path,
        through_stage="implicit_reasoning",
    )
    final_graph_path = Path(result["artifact_paths"]["final_graph_bundle"])
    output_root = tmp_path / "exports" / "import" / "示例法"
    split_graph_export(final_graph_path, output_root)

    assert (output_root / "neo4j" / "nodes.jsonl").exists()
    assert (output_root / "neo4j" / "edges.jsonl").exists()
    assert (output_root / "elasticsearch" / "documents.jsonl").exists()


def test_stage_bar_display_places_summary_under_progress_bar(capsys) -> None:
    display = StageBarDisplay()

    display.announce_discovery(2)
    display.start_stage("normalize")
    display.update(0, 3)
    display.update(1, 3)
    display.stage_summary("normalize", {"succeeded": 2, "failed": 1})
    display.start_stage("reference_filter:collect")
    display.update(2, 5)
    display.start_stage("relation_classify:model")
    display.update(1, 2)
    display.start_stage("structure")
    display.close()

    stderr = capsys.readouterr().err
    assert "law-kg build: 2 documents" in stderr
    assert "[normalize]:" in stderr
    assert "normalize" in stderr and "1/3" in stderr
    assert "2 succeed, 1 failed" in stderr
    assert "[reference_filter]:" in stderr
    assert "[relation_classify]:" in stderr
    assert "collect" in stderr
    assert "model" in stderr
    assert "[structure]:" in stderr


def test_builder_cli_accepts_multiple_categories() -> None:
    parser = build_parser()

    args = parser.parse_args(["build", "--category", "法律", "司法解释"])

    assert args.category == ["法律", "司法解释"]


def test_builder_cli_accepts_incremental_flag() -> None:
    parser = build_parser()

    args = parser.parse_args(["build", "--category", "法律", "--incremental"])

    assert args.incremental is True


def test_builder_cli_accepts_multiple_source_ids() -> None:
    parser = build_parser()

    args = parser.parse_args(["build", "--source-id", "law:one", "law:two"])

    assert args.source_ids == ["law:one", "law:two"]


def test_resolve_source_id_accepts_known_id(tmp_path: Path) -> None:
    source_id = write_pipeline_source(tmp_path)
    assert resolve_source_id(source_id, tmp_path) == source_id


def test_batch_build_runs_normalize_then_structure(tmp_path: Path) -> None:
    write_pipeline_source(tmp_path, source_id="law:sample")
    write_pipeline_source(tmp_path, source_id="interpretation:sample", title="司法解释示例")

    result = build_batch_knowledge_graph(
        data_root=tmp_path,
        through_stage="structure",
    )

    assert result["status"] == "completed"
    normalize_index = read_normalize_index(tmp_path / "intermediate" / "builder" / "01_normalize" / "normalize_index.json")
    assert normalize_index.stats["succeeded_sources"] == 2
    stage_graph = read_graph_bundle(tmp_path / "intermediate" / "builder" / "02_structure" / "graph_bundle-0001.json")
    document_nodes = [node for node in stage_graph.nodes if node.level == "document"]
    assert len(document_nodes) == 2
    assert {node.id for node in document_nodes} == {"law:sample", "interpretation:sample"}


def test_structure_supports_incremental_category_merge(tmp_path: Path) -> None:
    write_pipeline_source(tmp_path, source_id="law:sample", category="法律")
    write_pipeline_source(tmp_path, source_id="interpretation:sample", title="司法解释示例", category="司法解释", issuer="最高人民法院")
    build_batch_knowledge_graph(
        data_root=tmp_path,
        through_stage="normalize",
    )

    build_batch_knowledge_graph(
        data_root=tmp_path,
        category=["法律"],
        start_stage="structure",
        through_stage="structure",
    )
    first_graph = read_graph_bundle(tmp_path / "intermediate" / "builder" / "02_structure" / "graph_bundle-0001.json")
    first_document_ids = {node.id for node in first_graph.nodes if node.level == "document"}
    assert first_document_ids == {"law:sample"}

    build_batch_knowledge_graph(
        data_root=tmp_path,
        category=["司法解释"],
        start_stage="structure",
        through_stage="structure",
        incremental=True,
    )
    merged_graph = read_graph_bundle(tmp_path / "intermediate" / "builder" / "02_structure" / "graph_bundle-0001.json")
    merged_document_ids = {node.id for node in merged_graph.nodes if node.level == "document"}
    assert merged_document_ids == {"law:sample", "interpretation:sample"}
    assert set(merged_graph.metadata["processed_source_ids"]) == {"law:sample", "interpretation:sample"}


def test_reference_filter_and_relation_classify_can_process_selected_categories_incrementally(tmp_path: Path) -> None:
    write_pipeline_source(tmp_path, source_id="law:sample", category="法律")
    write_pipeline_source(
        tmp_path,
        source_id="interpretation:sample",
        title="关于示例法的解释",
        category="司法解释",
        issuer="最高人民法院",
    )
    build_batch_knowledge_graph(
        data_root=tmp_path,
        through_stage="structure",
    )

    build_batch_knowledge_graph(
        data_root=tmp_path,
        category=["法律"],
        start_stage="reference_filter",
        through_stage="relation_classify",
    )
    first_candidates = read_reference_candidates(tmp_path / "intermediate" / "builder" / "03_reference_filter" / "candidates.jsonl")
    first_plans = read_relation_plans(tmp_path / "intermediate" / "builder" / "04_relation_classify" / "results.jsonl")
    assert first_candidates
    assert first_plans
    assert all(row.id.startswith("law:") for row in first_candidates)
    assert all(row.label == "REFERENCES" for row in first_plans)
    assert not any(row.id.startswith("interpretation:") for row in first_plans)

    build_batch_knowledge_graph(
        data_root=tmp_path,
        category=["司法解释"],
        start_stage="reference_filter",
        through_stage="relation_classify",
        incremental=True,
    )
    merged_candidates = read_reference_candidates(tmp_path / "intermediate" / "builder" / "03_reference_filter" / "candidates.jsonl")
    merged_plans = read_relation_plans(tmp_path / "intermediate" / "builder" / "04_relation_classify" / "results.jsonl")
    assert any(row.id.startswith("interpretation:") for row in merged_candidates)
    assert any(row.id.startswith("interpretation:") for row in merged_plans)


def test_entity_extraction_supports_incremental_category_merge(tmp_path: Path) -> None:
    write_pipeline_source(tmp_path, source_id="law:sample", category="法律")
    write_pipeline_source(
        tmp_path,
        source_id="interpretation:sample",
        title="关于示例法的解释",
        category="司法解释",
        issuer="最高人民法院",
    )
    build_batch_knowledge_graph(
        data_root=tmp_path,
        through_stage="relation_classify",
    )

    build_batch_knowledge_graph(
        data_root=tmp_path,
        category=["法律"],
        start_stage="entity_extraction",
        through_stage="entity_extraction",
    )
    first_graph = read_graph_bundle(tmp_path / "intermediate" / "builder" / "05_entity_extraction" / "graph_bundle-0001.json")
    first_context = build_relation_context(first_graph)
    first_mentions = [edge for edge in first_graph.edges if edge.type == "MENTIONS"]
    assert first_mentions
    assert all(first_context["owner_document_by_node"].get(edge.source, edge.source) == "law:sample" for edge in first_mentions)

    build_batch_knowledge_graph(
        data_root=tmp_path,
        category=["司法解释"],
        start_stage="entity_extraction",
        through_stage="entity_extraction",
        incremental=True,
    )
    merged_graph = read_graph_bundle(tmp_path / "intermediate" / "builder" / "05_entity_extraction" / "graph_bundle-0001.json")
    merged_context = build_relation_context(merged_graph)
    merged_mentions = [edge for edge in merged_graph.edges if edge.type == "MENTIONS"]
    assert set(merged_graph.metadata["processed_source_ids"]) == {"law:sample", "interpretation:sample"}
    assert any(
        merged_context["owner_document_by_node"].get(edge.source, edge.source) == "interpretation:sample"
        for edge in merged_mentions
    )


def test_batch_build_continues_when_some_documents_are_missing(tmp_path: Path) -> None:
    write_pipeline_source(tmp_path, source_id="law:ok")
    metadata_path = tmp_path / "source" / "metadata" / "metadata-0002.json"
    metadata_path.write_text(
        json.dumps(
            [
                {
                    "source_id": "law:missing",
                    "title": "损坏文档",
                    "category": "法律",
                    "source_format": "docx",
                }
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    result = build_batch_knowledge_graph(
        data_root=tmp_path,
        through_stage="normalize",
    )

    assert result["status"] == "partial"
    assert result["completed_count"] == 1
    assert result["failed_count"] == 1
    normalize_index = read_normalize_index(tmp_path / "intermediate" / "builder" / "01_normalize" / "normalize_index.json")
    assert normalize_index.stats["succeeded_sources"] == 1

    manifest = read_manifest(Path(result["manifest_path"]))
    normalize_stage = next(stage for stage in manifest.stages if stage.name == "normalize")
    assert len(normalize_stage.failures) == 1
    assert normalize_stage.failures[0]["source_id"] == "law:missing"
    assert "Document not found" in normalize_stage.failures[0]["message"]


def test_batch_build_reuses_existing_stage_outputs_by_default(tmp_path: Path) -> None:
    write_pipeline_source(tmp_path, source_id="law:sample")
    first = build_batch_knowledge_graph(
        data_root=tmp_path,
        through_stage="structure",
    )
    second = build_batch_knowledge_graph(
        data_root=tmp_path,
        through_stage="structure",
    )

    assert first["status"] == "completed"
    assert second["status"] == "completed"

    manifest = read_manifest(Path(second["manifest_path"]))
    normalize_stage = next(stage for stage in manifest.stages if stage.name == "normalize")
    structure_stage = next(stage for stage in manifest.stages if stage.name == "structure")

    assert normalize_stage.stats["reused_sources"] == 1
    assert structure_stage.stats["reused_sources"] == 1
