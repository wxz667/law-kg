from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from builder.io import read_normalize_index, read_normalized_document
from builder.stages.normalize import run as run_normalize
from builder.stages.structure import run as run_structure


def write_docx(path: Path, paragraphs: list[str]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(str(path))
    return path


def write_docx_with_table(
    path: Path,
    paragraphs: list[str],
    rows: list[list[str]],
) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    table = document.add_table(rows=len(rows), cols=len(rows[0]) if rows else 0)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    document.save(str(path))
    return path


def write_metadata(root: Path, payloads: list[dict[str, object]], file_name: str = "metadata-0001.json") -> None:
    metadata_dir = root / "source" / "metadata"
    metadata_dir.mkdir(parents=True, exist_ok=True)
    (metadata_dir / file_name).write_text(
        json.dumps(payloads, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def test_normalize_writes_per_document_artifact_and_preserves_metadata(tmp_path: Path) -> None:
    write_docx(
        tmp_path / "source" / "docs" / "示例法.docx",
        [
            "示例法",
            "第一条 为了规范道路交通活动，制定本法。",
            "第二条 本法适用于道路交通活动。",
            "附件一",
            "附表内容",
        ],
    )
    write_metadata(
        tmp_path,
        [
            {
                "source_id": "law:sample",
                "title": "示例法",
                "category": "法律",
                "status": "现行有效",
                "issuer": "全国人民代表大会常务委员会",
                "publish_date": "2024-01-01",
                "effective_date": "2024-02-01",
                "source_url": "https://example.com/law",
                "source_format": "docx",
            }
        ],
    )

    index = run_normalize(tmp_path)
    artifact_path = tmp_path / "intermediate" / "builder" / "01_normalize" / "documents" / "law:sample.json"
    document = read_normalized_document(artifact_path)

    assert index.stats["succeeded_sources"] == 1
    assert artifact_path.exists()
    assert document.source_id == "law:sample"
    assert document.title == "示例法"
    assert "第一条 为了规范道路交通活动，制定本法。" in document.content
    assert document.appendix_lines == ["附件一", "附表内容"]
    assert document.metadata["category"] == "法律"
    assert "source_format" not in document.to_dict()


def test_normalize_matches_docx_by_whitespace_normalization(tmp_path: Path) -> None:
    write_docx(
        tmp_path / "source" / "docs" / "最高人民法院 国务院宗教事务局关于寺庙、道观房屋产权归属问题的复函.docx",
        [
            "最高人民法院 国务院宗教事务局关于寺庙、道观房屋产权归属问题的复函",
            "第一条 正文。",
        ],
    )
    write_metadata(
        tmp_path,
        [
            {
                "source_id": "reply:sample",
                "title": "最高人民法院　国务院宗教事务局关于寺庙、道观房屋产权归属问题的复函",
                "category": "司法解释",
                "source_format": "docx",
            }
        ],
    )

    index = run_normalize(tmp_path)
    entry = index.entries[0]

    assert entry.status == "completed"
    assert entry.document_path.endswith(".docx")


def test_normalize_records_missing_documents_without_stopping(tmp_path: Path) -> None:
    write_docx(
        tmp_path / "source" / "docs" / "存在的法.docx",
        [
            "存在的法",
            "第一条 正文。",
        ],
    )
    write_metadata(
        tmp_path,
        [
            {"source_id": "law:ok", "title": "存在的法", "category": "法律", "source_format": "docx"},
            {"source_id": "law:missing", "title": "不存在的法", "category": "法律", "source_format": "docx"},
        ],
    )

    index = run_normalize(tmp_path)
    saved_index = read_normalize_index(tmp_path / "intermediate" / "builder" / "01_normalize" / "normalize_index.json")
    log_payload = json.loads((tmp_path.parent / "logs" / "builder" / "normalize-report.json").read_text(encoding="utf-8"))

    assert index.stats["succeeded_sources"] == 1
    assert index.stats["failed_sources"] == 1
    assert saved_index.stats == index.stats
    missing_entry = next(entry for entry in index.entries if entry.source_id == "law:missing")
    assert missing_entry.error_type == "missing_document"
    assert any(entry["source_id"] == "law:missing" for entry in log_payload["entries"])


def test_normalize_keeps_only_embedded_substantive_document(tmp_path: Path) -> None:
    write_docx(
        tmp_path / "source" / "docs" / "通知型规定.docx",
        [
            "最高人民法院印发《关于审理行政赔偿案件若干问题的规定》的通知",
            "全国地方各级人民法院：",
            "现将《关于审理行政赔偿案件若干问题的规定》印发给你们，请遵照执行。",
            "附：",
            "最高人民法院",
            "关于审理行政赔偿案件若干问题的规定",
            "（1997年12月5日最高人民法院审判委员会第950次会议通过）",
            "第一条 人民法院审理行政赔偿案件，适用本规定。",
            "第二条 行政赔偿请求依法处理。",
        ],
    )
    write_metadata(
        tmp_path,
        [
            {
                "source_id": "notice:embedded",
                "title": "通知型规定",
                "category": "司法解释",
                "source_format": "docx",
            }
        ],
    )

    run_normalize(tmp_path)
    document = read_normalized_document(tmp_path / "intermediate" / "builder" / "01_normalize" / "documents" / "notice:embedded.json")

    assert document.title == "通知型规定"
    assert "第一条" in document.content


def test_normalize_flattens_table_rows_into_numbered_content(tmp_path: Path) -> None:
    write_docx_with_table(
        tmp_path / "source" / "docs" / "罪名对应表.docx",
        ["罪名对应表"],
        [
            ["刑法条文", "罪名"],
            ["第一百三十三条之一\n（《刑法修正案（八）》第二十二条）", "危险驾驶罪"],
            ["第一百四十三条\n（《刑法修正案（八）》第二十四条）", "生产、销售不符合安全标准的食品罪"],
        ],
    )
    write_metadata(
        tmp_path,
        [
            {
                "source_id": "law:crime-table",
                "title": "罪名对应表",
                "category": "法律",
                "source_format": "docx",
            }
        ],
    )

    run_normalize(tmp_path)
    document = read_normalized_document(tmp_path / "intermediate" / "builder" / "01_normalize" / "documents" / "law:crime-table.json")

    assert "1. 刑法条文：第一百三十三条之一 （《刑法修正案（八）》第二十二条）" in document.content
    assert "罪名：危险驾驶罪；" in document.content
    assert "2. 刑法条文：第一百四十三条 （《刑法修正案（八）》第二十四条）" in document.content


def test_structure_uses_source_id_and_parses_appendix(tmp_path: Path) -> None:
    write_docx(
        tmp_path / "source" / "docs" / "示例法.docx",
        [
            "示例法",
            "第一章 总则",
            "第一条 为了规范道路交通活动，制定本法。",
            "第二条 本法适用于道路交通活动。",
            "附件一",
            "第一项 附件事项",
        ],
    )
    write_metadata(
        tmp_path,
        [
            {
                "source_id": "law:sample",
                "title": "示例法",
                "category": "法律",
                "document_type": "law",
                "source_format": "docx",
            }
        ],
    )

    run_normalize(tmp_path)
    bundle = run_structure(tmp_path)

    document_node = next(node for node in bundle.nodes if node.level == "document")
    assert document_node.id == "law:sample"
    assert document_node.name == "示例法"
    assert document_node.category == "法律"
    assert any(node.level == "article" and node.name == "第一条" for node in bundle.nodes)
    assert any(node.level == "appendix" for node in bundle.nodes)


def test_structure_falls_back_to_single_body_segment_for_unstructured_content(tmp_path: Path) -> None:
    write_docx(
        tmp_path / "source" / "docs" / "批复示例.docx",
        [
            "最高人民检察院关于对报请批准逮捕的案件可否侦查问题的批复",
            "海南省人民检察院：",
            "你院请示收悉。经研究，批复如下：",
            "人民检察院审查公安机关提请逮捕的案件，经审查，应当作出批准或者不批准逮捕的决定。",
        ],
    )
    write_metadata(
        tmp_path,
        [
            {
                "source_id": "reply:sample",
                "title": "批复示例",
                "category": "司法解释",
                "document_type": "interpretation",
                "source_format": "docx",
            }
        ],
    )

    run_normalize(tmp_path)
    bundle = run_structure(tmp_path)
    segments = [node for node in bundle.nodes if node.level == "segment"]

    assert len(segments) == 1
    assert segments[0].name == "正文"
    assert segments[0].text.startswith("海南省人民检察院：")
    assert bundle.metadata["stage"] == "structure"


def test_normalize_filters_toc_and_structure_collapses_single_paragraph_items_to_article(tmp_path: Path) -> None:
    write_docx(
        tmp_path / "source" / "docs" / "示例规则.docx",
        [
            "示例规则",
            "目录",
            "第一百一十条第四款 有下列情形之一的，属于有碍侦查：",
            "第一百一十条第五款 可能毁灭、伪造证据，干扰证人作证或者串供的；",
            "第一百一十条第六款 可能自杀或者逃跑的；",
            "第一百一十条 有下列情形之一的，属于有碍侦查：",
            "（一）可能毁灭、伪造证据，干扰证人作证或者串供的；",
            "（二）可能自杀或者逃跑的；",
            "（三）可能导致同案犯逃避侦查的；",
        ],
    )
    write_metadata(
        tmp_path,
        [
            {
                "source_id": "rule:sample",
                "title": "示例规则",
                "category": "司法解释",
                "source_format": "docx",
            }
        ],
    )

    run_normalize(tmp_path, force_rebuild=True)
    document = read_normalized_document(tmp_path / "intermediate" / "builder" / "01_normalize" / "documents" / "rule:sample.json")
    assert "目录" not in document.content
    assert "第一百一十条第四款 有下列情形之一的，属于有碍侦查：" not in document.content

    bundle = run_structure(tmp_path)
    articles = [node for node in bundle.nodes if node.level == "article"]
    paragraphs = [node for node in bundle.nodes if node.level == "paragraph"]
    items = [node for node in bundle.nodes if node.level == "item"]

    assert any(node.name == "第一百一十条" and node.text == "有下列情形之一的，属于有碍侦查：" for node in articles)
    assert len([node for node in paragraphs if node.name.startswith("第一百一十条")]) == 0
    assert len(items) == 3
    assert any(node.name == "第一百一十条第一项" for node in items)
    assert any(node.text == "可能毁灭、伪造证据，干扰证人作证或者串供的；" for node in items)


def test_structure_treats_semicolon_clauses_after_colon_as_items(tmp_path: Path) -> None:
    write_docx(
        tmp_path / "source" / "docs" / "特别重大贿赂犯罪规则.docx",
        [
            "特别重大贿赂犯罪规则",
            "第四十五条 对于特别重大贿赂犯罪案件,犯罪嫌疑人被羁押或者监视居住的，应当经许可。",
            "有下列情形之一的，属于特别重大贿赂犯罪：",
            "涉嫌贿赂犯罪数额在五十万元以上，犯罪情节恶劣的；",
            "有重大社会影响的；",
            "涉及国家重大利益的。",
        ],
    )
    write_metadata(
        tmp_path,
        [
            {
                "source_id": "rule:semicolon-items",
                "title": "特别重大贿赂犯罪规则",
                "category": "司法解释",
                "source_format": "docx",
            }
        ],
    )

    run_normalize(tmp_path, force_rebuild=True)
    bundle = run_structure(tmp_path)

    paragraphs = [node for node in bundle.nodes if node.level == "paragraph"]
    items = [node for node in bundle.nodes if node.level == "item"]

    assert any(node.name == "第四十五条第二款" and node.text == "有下列情形之一的，属于特别重大贿赂犯罪：" for node in paragraphs)
    assert len(items) == 3
    assert [node.text for node in items] == [
        "涉嫌贿赂犯罪数额在五十万元以上，犯罪情节恶劣的；",
        "有重大社会影响的；",
        "涉及国家重大利益的。",
    ]
